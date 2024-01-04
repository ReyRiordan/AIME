from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
import time
import datetime as date
from docx import Document
import io
import os
import streamlit as st
import streamlit.components.v1 as components
import base64
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName,
    FileType, Disposition, ContentId)
from constants import *
from email import send_email
import descriptions


# SECRETS
os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
os.environ["LOGIN_PASS"] = st.secrets["LOGIN_PASS"]
LOGIN_PASS = os.getenv("LOGIN_PASS")


st.title("Medical Interview Simulation")

if "stage" not in st.session_state:
    st.session_state["stage"] = LOGIN_PAGE

def set_stage(stage):
    st.session_state["stage"] = stage


if st.session_state["stage"] == LOGIN_PAGE:
    st.write(descriptions.get("intro"))
    
    st.session_state["username"] = st.text_input("Enter any username (does not have to be your real name) and press \"Enter\":")
    if st.session_state["username"]:
        password = st.text_input("Enter the password you were provided and press \"Enter\":")
        if password == LOGIN_PASS: 
            st.write("Authentication successful!")
            time.sleep(2)
            set_stage(PATIENT_SELECTION)
            st.rerun()


if st.session_state["stage"] == PATIENT_SELECTION:
    st.write(descriptions.get("selection"))
    
    st.session_state["interview"] = None
    st.session_state["feedback_string"] = None
    st.session_state["messages"] = []
    st.session_state["patient"] = st.selectbox("Which patient would you like to interview?", 
                                               ["John Smith", "Jackie Smith"],
                                               index = None,
                                               placeholder = "Select patient...")

    st.button("Start Interview", on_click=set_stage, args=[PATIENT_LOADING])


if st.session_state["stage"] == PATIENT_LOADING:
    with open(BASE_PROMPT, 'r', encoding='utf8') as base:
        base_prompt = base.read()
    INFO = prompts[st.session_state["patient"]]
    with open(INFO, 'r', encoding='utf8') as info:
        patient_info = info.read()
    prompt_input = str(base_prompt + patient_info)

    MODEL = "gpt-4"
    llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model_name=MODEL, temperature=0.0)
    st.session_state["conversation"] = ConversationChain(llm=llm, memory=ConversationBufferMemory())
    initial_output = st.session_state["conversation"].predict(input=prompt_input)

    st.session_state["messages"].append({"role": "Assistant", "content": "You may now begin your interview with " + st.session_state["patient"] + "."})
    
    set_stage(CHAT_INTERFACE)


if st.session_state["stage"] == CHAT_INTERFACE:
    st.write(descriptions.get("interview"))
    
    for message in st.session_state["messages"]:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if user_input := st.chat_input("Type here..."):
        with st.chat_message(st.session_state["username"]):
            st.markdown(user_input)
        st.session_state["messages"].append({"role": st.session_state["username"], "content": user_input})
        output = st.session_state["conversation"].predict(input=user_input)
        with st.chat_message(st.session_state["patient"]):
            st.markdown(output)
            st.session_state["messages"].append({"role": st.session_state["patient"], "content": output})

    st.button("End Interview", on_click=set_stage, args=[CREATE_INTERVIEW_FILE])


if st.session_state["stage"] == CREATE_INTERVIEW_FILE:
    st.session_state["interview"] = Document()
    heading = st.session_state["interview"].add_paragraph("User: " + st.session_state["username"] + ", ")
    currentDateAndTime = date.datetime.now()
    date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")
    heading.add_run("Date: " + date_time + ", ")
    heading.add_run("Patient: " + st.session_state["patient"])
    for message in st.session_state["messages"]:
        st.session_state["interview"].add_paragraph(message["role"] + ": " + message["content"])
    st.session_state["interview"].save("./Conversations/" + st.session_state["username"]+"_"+date_time+".docx")

    set_stage(POST_INTERVIEW)


if st.session_state["stage"] == POST_INTERVIEW:
    
    st.write(descriptions.get("post"))
    
    st.button("View Physical", on_click=set_stage, args=[PHYSICAL_SCREEN])
    st.button("View ECG", on_click=set_stage, args=[ECG_SCREEN])
    st.button("Provide Your Feedback", on_click=set_stage, args=[FEEDBACK_SCREEN])


if st.session_state["stage"] == PHYSICAL_SCREEN:
    st.header("Physical Examination Findings")
    st.write("Here is the full physical examination for " + st.session_state["patient"] + ". Click the \"Back\" button to go back once you're done.")
    if st.session_state["patient"] == "John Smith":
        physical_exam_doc = Document(PHYSICAL_LOCATION_JOHN)
    else:
        physical_exam_doc = Document(PHYSICAL_LOCATION_JACKIE)
    
    for parargraph in physical_exam_doc.paragraphs:
        st.write(parargraph.text)
    st.button("Back", on_click=set_stage, args=[POST_INTERVIEW])
    

if st.session_state["stage"] == ECG_SCREEN:
    st.header("ECG Chart")
    st.write("Here is the ECG for " + st.session_state["patient"] + ". Click the \"Back\" button to go back once you're done.")
    if st.session_state["patient"] == "John Smith":
        st.image(ECG_LOCATION_JOHN)
    else:
        st.image(ECG_LOCATION_JACKIE)
    st.button("Back", on_click=set_stage, args=[POST_INTERVIEW])


if st.session_state["stage"]==FEEDBACK_SCREEN:
    st.write(descriptions.get("feedback"))
    st.session_state["diagnosis"]=st.text_area("Provide your differential diagnosis here:")
    st.session_state["feedback_string"] = st.text_area("Provide your feedback here:")

    st.button("Send", on_click=set_stage, args=[FINAL_SCREEN])
    st.button("Back", on_click=set_stage, args=[POST_INTERVIEW])
    
    for message in st.session_state["messages"]:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    st.session_state["feedback_string"] = "<p> " + st.session_state["diagnosis"]+"</p> <p> "+st.session_state["feedback_string"]+" </p>"
    st.session_state["feedback_string"] = "<h2>User: "+st.session_state["username"]+ "</h2> <h3> Feedback: </h3>" + st.session_state["feedback_string"]


if st.session_state["stage"] == FINAL_SCREEN: 
    st.write(descriptions.get("final"))

    currentDateAndTime = date.datetime.now()
    date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")

    # Setting up file for attachment sending
    bio = io.BytesIO()
    st.session_state["interview"].save(bio)
    send_email(bio, EMAIL_TO_SEND, st.session_state["username"], date_time, st.session_state["feedback_string"])
    
    st.download_button("Download interview", 
                        data=bio.getvalue(),
                        file_name=st.session_state["username"]+"_"+date_time+".docx",
                        mime="docx")
    
    st.button("New Interview", on_click=set_stage, args=[PATIENT_SELECTION])
    
    for message in st.session_state["messages"]:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])