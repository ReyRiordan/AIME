import time
from datetime import datetime
from docx import Document
import io
import os
import streamlit as st
import streamlit.components.v1 as components
# import streamlit_authenticator as auth
import base64
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName,
    FileType, Disposition, ContentId)
from audiorecorder import audiorecorder
from openai import OpenAI
import tempfile
from annotated_text import annotated_text
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from lookups import *
from web_classes import *
from web_methods import *
import pytz


# STREAMLIT SETUP
st.set_page_config(page_title = "MEWAI",
                   page_icon = "ðŸ§‘â€âš•ï¸",
                   layout = "wide",
                   initial_sidebar_state="collapsed")

if "stage" not in st.session_state:
    st.session_state["stage"] = LOGIN_PAGE

def set_stage(stage):
    st.session_state["stage"] = stage

def update_interview():
    COLLECTION.replace_one({"username" : st.session_state["username"], 
                            "start_time" : st.session_state["start_time"]}, 
                            st.session_state["interview"].model_dump())
    
def read_time(iso_time) -> str:
    if not iso_time: return "N/A"
    dt = datetime.fromisoformat(iso_time)
    est = pytz.timezone("US/Eastern")
    dt_est = dt.astimezone(est)
    return dt_est.strftime("%B %d, %Y at %I:%M %p")


# DB SETUP
@st.cache_resource
def init_connection():
    return MongoClient(DB_URI)

DB_CLIENT = init_connection()
COLLECTION = DB_CLIENT[DB_NAME]["Interviews"]

def get_data(username: str = None) -> list[dict]:
    DB = DB_CLIENT[DB_NAME]
    items = DB["Interviews"].find()
    items = list(items)  # make hashable for st.cache_data

    if username:
        only_user = []
        for item in items:
            if item["username"] == username: only_user.append(item)
        items = only_user
    
    items = sorted(items, key=lambda x: datetime.fromisoformat(x["start_time"]), reverse=True)
    return items


# APP CODE STARTS HERE

if st.session_state["stage"] == ERROR:
    layout1 = st.columns([2, 3, 2])
    with layout1[1]:
        st.title("ERROR")
        st.write("An error has occurred. No worries - the admin has been alerted, and your responses so far have been automatically saved.")
        st.write(f"Error details: {st.session_state['error_details']}")


if st.session_state["stage"] == CLOSED:
    layout1 = st.columns([2, 3, 2])
    with layout1[1]:
        st.title("CURRENTLY CLOSED")
        st.write("We are experiencing an outage beyond our control that is delaying the release - please come back later :(")
        st.write("Rest assured, the deadline will be extended if this issue persists.")
        st.write("rhr58@scarletmail.rutgers.edu")


if st.session_state["stage"] == LOGIN_PAGE:
    layout1 = st.columns([2, 3, 2])
    with layout1[1]:
        st.title("Medical Interview Simulation (BETA)")
        st.write("Welcome! This is a WIP application where you can interview AI patients, write a post note, and automatically receive feedback on your performance.")
        st.write("Begin by logging in as directed. If you encounter any issues, please contact rhr58@scarletmail.rutgers.edu")

        username = st.text_input("Username (NetID):")
        if username and username not in ASSIGNMENTS: 
            st.write("Invalid username.")
        st.session_state["admin"] = True if username == "admin" else False
        password = st.text_input("Password (FirstLast):", type = "password")

        layout12b = layout1[1].columns(5)
        if layout12b[2].button("Log in"):
            if username in ASSIGNMENTS:
                correct = ASSIGNMENTS[username]["First_name"] + ASSIGNMENTS[username]["Last_name"]
                if username in ASSIGNMENTS and (password == correct or password == username):
                    st.session_state["username"] = username
                    st.session_state["assignment"] = ASSIGNMENTS[username]
                    st.write("Authentication successful!")
                    time.sleep(1)
                    set_stage(SETTINGS)
                    st.rerun()
                else:
                    st.write("Password incorrect.")


if st.session_state["stage"] == SETTINGS:
    st.session_state["interview"] = None
    st.session_state["messages"] = []
    st.session_state["convo_memory"] = []
    st.session_state["convo_summary"] = ""
    st.session_state["convo_prompt"] = ""
    st.session_state["start_time"] = datetime.now().isoformat()
    st.session_state["tokens"] = {"convo": {"input": 0, "output": 0},
                                  "feedback": {"input": 0, "output": 0}}
    st.session_state["saved_inputs"] = {"HPI": "", 
                                        "Past Histories": "",
                                        "Summary Statement": "", 
                                        "Assessment": "", 
                                        "Plan": ""}

    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        if st.session_state["admin"]:
            st.title("Patient Settings")
            patient_name = st.selectbox("Which patient would you like to interview?", 
                                        ["Jeffrey Smith", "Jenny Smith", "Samuel Thompson", "Sarah Thompson"],
                                        index = None,
                                        placeholder = "Select patient...")
        else:
            st.title("Patient Settings")
            DATA = get_data(st.session_state["username"])
            interview_list = {"First case": -99, "Second case": -99}
            for i in range(len(DATA)):
                interview = DATA[i]
                if interview["finished"] == False:
                    label = "CONTINUE: " + interview["patient"]["id"] + " @ " + read_time(interview["start_time"])
                    interview_list[label] = i

            case = st.selectbox("Are you doing your first or second case? Please make sure to do both. You may also continue a previous unfinished case.", 
                                    options = interview_list, 
                                    index = None, 
                                    placeholder = "Select case...")
            patient_name = None
            continue_previous = False

            if case:
                if case in ["First case", "Second case"]:
                    case_number = case.replace(" ", "_")
                    gender = st.session_state["assignment"][case_number]
                    if case_number == "First_case":
                        if gender == "M": patient_name = "Jeffrey Smith"
                        elif gender == "F": patient_name = "Jenny Smith"
                    elif case_number == "Second_case":
                        if gender == "M": patient_name = "Samuel Thompson"
                        elif gender == "F": patient_name = "Sarah Thompson"
                else:
                    case_index = interview_list[case]
                    continue_previous = True

        chat_mode = st.selectbox("Would you like to use text or voice input for the interview? Voice is encouraged for both practice and testing purposes.",
                                ["Text", "Voice"],
                                index = None,
                                placeholder = "Select interview mode...")
        
        # ADD ASSIGNMENT INFO?
        if st.button("Start Interview"):
            if continue_previous and chat_mode:
                TBC = DATA[case_index]
                st.session_state["interview"] = Interview.restore_previous(TBC)
                st.session_state["start_time"] = st.session_state["interview"].start_time
                st.session_state["convo_prompt"] = st.session_state["interview"].patient.convo_prompt
                st.session_state["interview"].record_time("continue")
                jump = 0 # WHERE TO CONTINUE FROM
                if chat_mode == "Text": jump = CHAT_INTERFACE_TEXT
                elif chat_mode == "Voice": jump = CHAT_INTERFACE_VOICE
                if TBC["messages"]:
                    for message in TBC["messages"]:
                        st.session_state["interview"].add_message(Message(type=message["type"], role=message["role"], content=message["content"]))
                    if TBC["convo_data"]:
                        st.session_state["messages"] = TBC["convo_data"]["messages"]
                        st.session_state["convo_memory"] = TBC["convo_data"]["convo_memory"]
                        st.session_state["convo_summary"] = TBC["convo_data"]["convo_summary"]
                if TBC["post_note_inputs"]:
                    st.session_state["interview"].post_note_inputs = TBC["post_note_inputs"]
                    st.session_state["saved_inputs"] = TBC["post_note_inputs"]
                    jump = DIAGNOSIS
                if TBC["feedback"]:
                    st.session_state["interview"].feedback = Feedback.restore_previous(TBC["feedback"]["feedback"]) # handle that weird feedback nesting sht
                    jump = FEEDBACK_SCREEN
                set_stage(jump)
                st.rerun()

            elif patient_name and chat_mode:
                st.session_state["interview"] = Interview.build(username = st.session_state["username"], 
                                                                patient = Patient.build(patient_name), 
                                                                start_time = st.session_state["start_time"], 
                                                                chat_mode = chat_mode)
                st.session_state["interview"].record_time("start")
                st.session_state["convo_prompt"] = st.session_state["interview"].patient.convo_prompt
                COLLECTION.insert_one(st.session_state["interview"].model_dump()) # INSERT INTERVIEW

                if chat_mode == "Text": chat_page = CHAT_INTERFACE_TEXT
                elif chat_mode == "Voice": chat_page = CHAT_INTERFACE_VOICE
                set_stage(chat_page)
                st.rerun()

            else: st.write("Incomplete settings.")


if st.session_state["stage"] == CHAT_INTERFACE_VOICE:
    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        st.title("Interview")
        st.write("You may now begin your interview - start by introducing yourself. Click \"Save\" to save your progress.")
        st.write("Click the \"Start Recording\" button to start recording your voice message. The button will then turn into a Stop button, which you can click when you are done talking.")
        st.write("Once you decide that you're done interviewing, click \"End Interview\" to proceed.")

        audio = audiorecorder("Start Recording", "Stop")
        
        # Check if a new recording has started
        if len(audio) > 0:
            if st.session_state.get("last_audio") != audio:
                st.session_state["audio_handled"] = False
                st.session_state["last_audio"] = audio

        container = st.container(height=300)
        for message in st.session_state["interview"].messages:
            with container:
                with st.chat_message(message.role):
                    st.markdown(message.content)

        if len(audio) > 0 and not st.session_state["audio_handled"]:
            user_input = transcribe_voice(audio)
            with container:
                with st.chat_message("User"):
                    st.markdown(user_input)
            response, speech = get_chat_output(user_input)
            with container:
                with st.chat_message("AI"):
                    st.markdown(response)
                    play_voice(speech)
            st.session_state["audio_handled"] = True

        def next_stage(screen):
            st.session_state["interview"].update_tokens(st.session_state["tokens"])
            st.session_state["interview"].record_time("end_interview")
            update_interview()
            set_stage(screen)

        def save():
            st.session_state["interview"].store_convo_data()
            st.session_state["interview"].record_time("save_interview")
            update_interview()

        columns = st.columns(5)
        columns[1].button("Restart", on_click=next_stage, args=[SETTINGS])
        columns[2].button("Save", on_click=save)
        columns[3].button("End Interview", on_click=next_stage, args=[PHYSICAL_ECG_SCREEN])


if st.session_state["stage"] == CHAT_INTERFACE_TEXT:
    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        st.title("Interview")
        st.write("You may now begin your interview - start by introducing yourself. Click \"Save\" to save your progress.")
        st.write("Enter your message into the text box below the chat screen and either click \"Enter\" on your keyboard or the paper airplane button to send it.")
        st.write("Once you decide that you're done interviewing, click \"End Interview\" to proceed.")

        container = st.container(height=300)
        for message in st.session_state["interview"].messages:
            with container:
                with st.chat_message(message.role):
                    st.markdown(message.content)

        if user_input := st.chat_input("Type here..."):
            with container:
                with st.chat_message("User"):
                    st.markdown(user_input)
            response, speech = get_chat_output(user_input)
            with container:
                with st.chat_message("AI"): #TODO Needs avatar eventually
                    st.markdown(response)
                    play_voice(speech)

        def next_stage(screen):
            st.session_state["interview"].update_tokens(st.session_state["tokens"])
            st.session_state["interview"].record_time("end_interview")
            update_interview()
            set_stage(screen)

        def save():
            st.session_state["interview"].store_convo_data()
            st.session_state["interview"].record_time("save_interview")
            update_interview()

        columns = st.columns(5)
        columns[1].button("Restart", on_click=next_stage, args=[SETTINGS])
        columns[2].button("Save", on_click=save)
        columns[3].button("End Interview", on_click=next_stage, args=[PHYSICAL_ECG_SCREEN])


# if st.session_state["stage"] == KEY_PHYSICALS:
#     layout1 = st.columns([2, 3, 2])
#     with layout1[1]:
#         st.title("Key Physical Exam Information")
#         findings = st.text_area(label = "Based on your interview with the patient, state the key physical examination findings that you would look for. You will be provided the actual exam results on the next screen.",
#                                 height = 200)

#         layout11 = st.columns([1, 1, 1])
#         if layout11[1].button("Next", use_container_width=True):
#             st.session_state["interview"].add_key_findings(findings)
#             set_stage(PHYSICAL_ECG_SCREEN)
#             st.rerun()


if st.session_state["stage"] == PHYSICAL_ECG_SCREEN:
    
    layout1 = st.columns([1, 3, 1])
    with layout1[1].container():
        st.title("Physical Examination")
        st.write("This is the physical examination obtained for the patient you interviewed. You will still be able to view it later, but please review it before proceeding.")
        
        st.divider()
        physical_exam_doc = Document(st.session_state["interview"].patient.physical)
        for paragraph in physical_exam_doc.paragraphs:
            st.write(paragraph.text)
        st.divider()

        layout11 = st.columns([1, 1, 1])
        layout11[1].button("Proceed to Post Note", on_click=set_stage, args = [DIAGNOSIS], use_container_width=True)


if st.session_state["stage"] == DIAGNOSIS:
    st.write("Write your post note as directed, then click \"Get Feedback\" to see how you did. Note that you can click \"Save\" to save your work and expand the text boxes by dragging the bottom-right corner.")
    st.divider()

    # 2 column full width layout
    layout1 = st.columns([1, 1])

    # User inputs
    summary = layout1[0].text_area(label = "**Summary Statement:** Provide a concise summary statement that uses semantic vocabulary to highlight the most important elements from history and exam to interpret and represent the patientâ€™s main problem.", 
                                   height = 200, 
                                   value = st.session_state["saved_inputs"]["Summary Statement"])
    assessment = layout1[0].text_area(label = "**Assessment**: Provide a differential diagnosis and explain the reasoning behind each diagnosis.", 
                                      height = 200, 
                                      value = st.session_state["saved_inputs"]["Assessment"])
    plan = layout1[0].text_area(label = "**Plan**: Include a diagnostic plan that explains the rationale for your decision. ", 
                                height = 200, 
                                value = st.session_state["saved_inputs"]["Plan"])

    # Interview transcription
    layout1[1].write("**Transcript**:")
    chat_container = layout1[1].container(height=400)
    for message in st.session_state["interview"].messages:
        with chat_container:
            with st.chat_message(message.role):
                st.markdown(message.content)
    # Physical Examination
    with layout1[1].expander("**Physical Examination**"):
        physical_exam_doc = Document(st.session_state["interview"].patient.physical)
        for paragraph in physical_exam_doc.paragraphs:
            st.write(paragraph.text)

    # 3 buttons: Get Feedback, New Interview, Download Interview
    st.divider()
    layout2 = st.columns([1, 1, 1, 1, 1])

    # Save
    if layout2[1].button("Save", use_container_width=True): 
        st.session_state["interview"].add_other_inputs("", "", summary, assessment, plan)
        st.session_state["saved_inputs"] = st.session_state["interview"].post_note_inputs
        st.session_state["interview"].record_time("save_postnote")
        update_interview()

    # Get Feedback
    if layout2[2].button("Get Feedback", use_container_width=True): 
        st.session_state["interview"].add_other_inputs("", "", summary, assessment, plan)
        st.session_state["saved_inputs"] = st.session_state["interview"].post_note_inputs
        st.session_state["interview"].record_time("get_feedback")
        update_interview()
        set_stage(FEEDBACK_SETUP)
        st.rerun()

    # New Interview
    # layout2[2].button("New Interview", on_click=set_stage, args=[SETTINGS], use_container_width=True)
    
    # Test cases
    if st.session_state["admin"]:
        layout21 = layout2[0].columns([1, 1])
        if layout21[0].button("TEST: BAD", use_container_width=True):
            with open("./IDEA/test_cases/bad.json", "r", encoding="utf8") as bad_json:
                bad_case = json.load(bad_json)
                # print(bad_case)
                # print("\n\n")
                st.session_state["interview"].add_key_findings(bad_case["Key Findings"])
                st.session_state["interview"].add_other_inputs("",
                                                            "",
                                                            bad_case["Summary"], 
                                                            bad_case["Assessment"], 
                                                            bad_case["Plan"])
                # print(st.session_state["interview"].post_note_inputs)
                # print("\n\n")
            update_interview()
            set_stage(FEEDBACK_SETUP)
            st.rerun()
        if layout21[1].button("TEST: GOOD", use_container_width=True):
            with open("./IDEA/test_cases/good.json", "r", encoding="utf8") as good_json:
                good_case = json.load(good_json)
                st.session_state["interview"].add_key_findings(good_case["Key Findings"])
                st.session_state["interview"].add_other_inputs("",
                                                            "",
                                                            good_case["Summary"], 
                                                            good_case["Assessment"], 
                                                            good_case["Plan"])
            update_interview()
            set_stage(FEEDBACK_SETUP)
            st.rerun()
    else:
        bio = io.BytesIO()
        st.session_state["convo_file"] = create_convo_file(st.session_state["interview"].username, 
                                                           st.session_state["interview"].patient.id, 
                                                           [message.model_dump() for message in st.session_state["interview"].messages])
        st.session_state["convo_file"].save(bio)

        layout2[3].download_button("Download interview", 
                                    data = bio.getvalue(), 
                                    file_name = st.session_state["interview"].username + "_" + st.session_state["start_time"] + ".docx", 
                                    mime = "docx")


if st.session_state["stage"] == FEEDBACK_SETUP:
    st.title("Processing feedback...")
    st.write("This might take a few minutes.")
    st.session_state["interview"].add_feedback(short=True)
    st.session_state["interview"].update_tokens(st.session_state["tokens"])
    st.session_state["interview"].record_time("feedback_processed")
    update_interview()
    set_stage(FEEDBACK_SCREEN)
    st.rerun()


if st.session_state["stage"] == FEEDBACK_SCREEN:
    st.title("Feedback")
    layout1 = st.columns([3, 1])
    layout1[0].write("The \"Post Note\" tab shows personalized feedback for each of your write-ups based on a detailed IDEA-based rubric. The \"Interview\" tab shows your interview transcript. The \"Case Explanation\" tab allows you to download a document with additional details and explanations on the patient case.")
    layout1[0].write("You're almost done! Click \"Next\" to proceed to the final screen.")
    layout11 = layout1[1].columns([1, 2, 1])
    layout11[1].button("**Next**", on_click=set_stage, args=[SURVEY], use_container_width=True, key=1)
    
    # Let the display methods cook
    display_Interview(st.session_state["interview"].model_dump())

    st.divider()
    layout2 = st.columns([1, 2, 1])
    layout2[1].button("**Next**", on_click=set_stage, args=[SURVEY], use_container_width=True, key=2)


if st.session_state["stage"] == SURVEY:
    layout1 = st.columns([2, 3, 2])
    with layout1[1]:
        st.title("Survey")
        response = st.text_area("Any feedback about your experience or suggestions for improvement?")
        if st.button("Finish"):
            if response:
                st.session_state["interview"].add_survey(response)
                st.session_state["interview"].record_time("end")
                st.session_state["interview"].finish()
                update_interview()
            set_stage(FINAL_SCREEN)
            st.rerun()


if st.session_state["stage"] == FINAL_SCREEN:
    layout1 = st.columns([2, 2, 2])
    with layout1[1]:
        st.title("Thank you! :heart:")
        st.write("Your responses have been saved automatically.")
        button_columns = st.columns(2)
        button_columns[0].button("New Interview", on_click=set_stage, args=[SETTINGS])
        button_columns[1].button("Back to Login", on_click=set_stage, args=[LOGIN_PAGE])