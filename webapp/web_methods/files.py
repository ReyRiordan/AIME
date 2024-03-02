from langchain_community.chat_models import ChatOpenAI
from langchain.chains.conversation.base import ConversationChain
from langchain.memory.buffer import ConversationBufferMemory
from docx import Document
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName,
    FileType, Disposition, ContentId)
import os
import datetime as date
import base64
import io
import streamlit as st
from audiorecorder import audiorecorder
from openai import OpenAI
import tempfile
from lookups import *
from web_classes import *
from annotated_text import annotated_text
import json


def create_convo_file(interview: Interview) -> Document:
    convo = Document()
    heading = convo.add_paragraph("User: " + interview.get_username() + ", ")
    currentDateAndTime = date.datetime.now()
    date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")
    heading.add_run("Date: " + date_time + ", ")
    heading.add_run("Patient: " + interview.get_patient().name)
    for message in interview.get_messages():
        convo.add_paragraph(message.role + ": " + message.content)
    
    return convo


def send_email(bio, EMAIL_TO_SEND, username, date_time, feedback):
    message = Mail(
        from_email = 'rutgers.aime@gmail.com',
        to_emails = EMAIL_TO_SEND,
        subject = "Conversation from " + username + " at time " + date_time,
        html_content = feedback)
    attachment = Attachment()
    encoded = base64.b64encode(bio.getvalue()).decode()
    attachment.file_content=FileContent(encoded)
    attachment.file_type = FileType('docx')
    attachment.file_name = FileName(username + "_" + date_time + ".docx")
    attachment.disposition = Disposition('attachment')
    attachment.content_id = ContentId('docx')
    message.attachment = attachment
    try:
        sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
        response = sg.send(message)
        print(response.status_code)
        print(response.body)
        print(response.headers)
    except: 
        print("ERROR ENCOUNTERED SENDING MESSAGE\n")