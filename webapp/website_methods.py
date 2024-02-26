from langchain_community.chat_models import ChatOpenAI
from langchain.chains.conversation.base import ConversationChain
from langchain.memory.buffer import ConversationBufferMemory
from docx import Document
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName,
    FileType, Disposition, ContentId)
import os
import datetime as date
import base64
import io
import streamlit as st
from audiorecorder import audiorecorder
from openai import OpenAI
import tempfile
from lookups import *
from website_classes import *
from annotated_text import annotated_text
import json


def get_webtext(content: str) -> str:
    path = WEBSITE_TEXT[content]
    with open(path, 'r', encoding='utf8') as webtext:
            text = webtext.read()
    return text


def create_convo_file(interview: Interview) -> Document:
    convo = Document()
    heading = convo.add_paragraph("User: " + interview.get_username() + ", ")
    currentDateAndTime = date.datetime.now()
    date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")
    heading.add_run("Date: " + date_time + ", ")
    heading.add_run("Patient: " + interview.get_patient().name)
    for message in interview.get_messages():
        convo.add_paragraph(message.role + ": " + message.content)
    
    return convo


def send_email(bio, EMAIL_TO_SEND, username, date_time, feedback):
    message = Mail(
        from_email = 'rutgers.aime@gmail.com',
        to_emails = EMAIL_TO_SEND,
        subject = "Conversation from " + username + " at time " + date_time,
        html_content = feedback)
    attachment = Attachment()
    encoded = base64.b64encode(bio.getvalue()).decode()
    attachment.file_content=FileContent(encoded)
    attachment.file_type = FileType('docx')
    attachment.file_name = FileName(username + "_" + date_time + ".docx")
    attachment.disposition = Disposition('attachment')
    attachment.content_id = ContentId('docx')
    message.attachment = attachment
    try:
        sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
        response = sg.send(message)
        print(response.status_code)
        print(response.body)
        print(response.headers)
    except: 
        print("ERROR ENCOUNTERED SENDING MESSAGE\n")


def transcribe_voice(voice_input, OPENAI_API_KEY):
    client = OpenAI()
    temp_audio_file = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
    voice_input.export(temp_audio_file.name, format="wav")
    with open(temp_audio_file.name, "rb") as file:
        transcription = client.audio.transcriptions.create(model="whisper-1", 
                                                    file=file, 
                                                    response_format="text")
    
    return transcription


def classifier(category: Category, messages: list[Message], OPENAI_API_KEY: str) -> None:
    # Create GPT instance
    llm = ChatOpenAI(api_key=OPENAI_API_KEY, model=CLASS_MODEL, temperature=0.0)
    conversation = ConversationChain(llm=llm, memory=ConversationBufferMemory())
    
    # Get the base class prompt for the category
    prompt_input = category.class_prompt

    # Append applicable messages to prompt
    applicable_messages = []
    for message in messages:
        if message.type == category.type:
            applicable_messages.append(message)
            message_content = message.content.rstrip() + "||"
            prompt_input += message_content

    # Classify
    raw_classification = conversation.predict(input=prompt_input)
    classifications = raw_classification.split("||")[:-1] # Remove empty classification at the end
    all_classifications = [] # list[list[str]]
    for classification in classifications:
        class_list = classification.split() # Split into individual labels by spaces
        for i in range(len(class_list)):
            class_list[i] = class_list[i].replace("_", " ") # Remove underscores from labels
        class_list = [label for label in class_list if label != "Other"] # Remove "Other" labels
        all_classifications.append(class_list) # Appends list of labels (one for each message)
    
    # Assign labels to each message accordingly
    if len(applicable_messages) != len(all_classifications):
        raise ValueError("Number of classifications must match number of applicable messages.")
    for i in range(len(applicable_messages)):
        if all_classifications[i]: # If not an empty list with no labels
            applicable_messages[i].labels[category.name] = all_classifications[i]
        

def annotate(interview: Interview, OPENAI_API_KEY: str) -> None:
    categories = interview.get_categories()
    messages = interview.get_messages()
    # Classify all messages
    for category in categories:
        classifier(category, messages, OPENAI_API_KEY)
    
    # Add highlight and annotation after all labels are assigned
    for message in messages:
        message.add_highlight()
        message.add_annotation()


def display_datagrades(grades: DataGrades, category: Category) -> None:
    score = grades.scores[category.name]
    max_score = grades.max_scores[category.name]
    st.header(f":{category.color}[{category.header}]: {score}/{max_score}", divider=category.color)
    display_labels = [(key, str(grades.weights[category.name][key]), "#baffc9" if value else "#ffb3ba") for key, value in grades.label_checklist[category.name].items()]
    annotated_text(display_labels)


def summarizer(LLM: OpenAI, convo_memory: list[dict[str, str]]) -> str:
    messages = [{"role": "system", "content": SUM_PROMPT}]
    dialogue = ""
    for message in convo_memory[1:]:
        if message["role"] == "system":
            dialogue += message["content"] + " \n"
        elif message["role"] == "user":
            dialogue += "User: " + message["content"] + " \n"
        elif message["role"] == "assistant":
            dialogue += "AI: " + message["content"] + "\n"
    messages.append({"role": "user", "content": dialogue})
    raw_summary = LLM.chat.completions.create(model = SUM_MODEL, 
                                              temperature = SUM_TEMP, 
                                              messages = messages)
    summary = raw_summary.choices[0].message.content
    return summary


def get_chat_output(LLM: OpenAI, convo_memory: list[dict[str, str]], user_input: str) -> list[list[dict[str, str]], str]:
    convo_memory.append({"role": "user", "content": user_input})
    response = LLM.chat.completions.create(model = CONVO_MODEL, 
                                           temperature = CHAT_TEMP, 
                                           messages = convo_memory)
    output = response.choices[0].message.content
    convo_memory.append({"role": "assistant", "content": output})
    if len(convo_memory) >= 10:
        summary = summarizer(LLM, convo_memory)
        convo_memory = [convo_memory[0], {"role": "system", "content": ("Summary of conversation so far: \n" + summary)}]
    return convo_memory, output


# TODO: @Rey since you have an aesthetically pleasing way of doing it please display all of the data from any given Interview instance. This method should take in an Interview and display all relavant data

def display_interview(interview: Interview):
    st.subheader("Username: " + interview.get_username())
    chat_container = st.container(height=300)
    for message in interview.get_messages():
            with chat_container:
                with st.chat_message(message.role):
                    if message.annotation is None:
                        st.markdown(message.content)
                    else:
                        annotated_text((message.content, message.annotation, message.highlight))
    # for category in interview.get_categories():
    #         if category.tab == "data":
    #             display_datagrades(interview.get_datagrades(), category)

def dict_to_patient(json_dict):
    to_return=Patient(json_dict["name"])
    return to_return

def dict_to_interview(json_dict):
    to_return=Interview(json_dict["username"],dict_to_patient(json_dict["patient"]))
    for message_json in json_dict["messages"]:
        to_return.add_message(Message(type=message_json["type"],
                                      role=message_json["role"],
                                      content=message_json["content"]
                                      ))
    return to_return