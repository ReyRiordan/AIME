import time
import datetime as date
from docx import Document
import io
import os
import streamlit as st
import streamlit.components.v1 as components
# import streamlit_authenticator as auth
import base64
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName,
    FileType, Disposition, ContentId)
from audiorecorder import audiorecorder
from openai import OpenAI
import tempfile
from annotated_text import annotated_text
from pymongo.mongo_client import MongoClient
from pymongo.server_api import ServerApi
from lookups import *
from web_classes import *
from web_methods import *

# from dotenv import load_dotenv

# load_dotenv()


############## DATABASE

# Establish connection to server

client = MongoClient(DB_URI,server_api=ServerApi('1'))

# Ping server on startup

try:
    client.admin.command('ping')
    print("Connection Successful")
except Exception as e:
    print(e)

# Method to get data of server

@st.cache_data(ttl=1200)
def get_data():
    db=client["AIME"]
    items=db["Conversation"].find()
    items = list(items)
    return items

# MongoDB Collection to add to

collection=client["AIME"]["Conversation"]

# Collection for surveys for FIRST BETATEST

survey = {}

######### WEBSITE #########

st.set_page_config(page_title = "AIME", page_icon = "🧑‍⚕️", layout = "wide")

if "stage" not in st.session_state:
    st.session_state["stage"] = LOGIN_PAGE

def set_stage(stage):
    st.session_state["stage"] = stage


if st.session_state["stage"] == LOGIN_PAGE:
    st.write("Welcome, and thank you for volunteering to participate in this beta test! This is an application where you will virtually simulate an interview with a patient, provide a differential diagnosis for them, and then automatically receive grading and feedback based on your performance.")
    st.write("Please follow the directions on each page to work through the whole application, and take notes on where there is potential room for improvement.")
    st.write("Begin by logging in!")
    
    layout1 = st.columns([2, 3, 2])
    with layout1[1]:
        st.title("Virtual Patient (BETA)")
        st.write("For beta testing use only.")

        username = st.text_input("Enter any username (full name):")
        password = st.text_input("Enter the password you were provided:", type = "password")

        layout12b = layout1[1].columns(5)
        if layout12b[2].button("Log in"):
            if username and password == LOGIN_PASS:
                st.session_state["username"] = username
                st.write("Authentication successful!")
                time.sleep(1)
                set_stage(SETTINGS)
                st.rerun()
            if username == DATABASE_USERNAME and password == DATABASE_PASSWORD:
                st.write("Authentication successful!")
                time.sleep(1)
                set_stage(VIEW_INTERVIEWS)
                st.rerun()
            else:
                st.write("Password incorect.")


if st.session_state["stage"]==VIEW_INTERVIEWS:
    st.title("View Interviews")
    if "interview_display_index" not in st.session_state:
        st.session_state["interview_display_index"]=0
        st.session_state["all_interviews"] = get_data() 



    list_of_interviews = {}
    for i in range(len(st.session_state["all_interviews"])):
        str_to_append = st.session_state["all_interviews"][i]["username"]+": "
        if "date_time" in st.session_state["all_interviews"][i].keys():
            str_to_append+=st.session_state["all_interviews"][i]["date_time"]
        list_of_interviews[str_to_append] = i
    

    interview_selection=st.selectbox("Select an interview", 
                                     options = list_of_interviews, 
                                     placeholder = "Select Interview")
    st.session_state["interview_display_index"] = list_of_interviews[interview_selection]

    st.subheader("Interview " + str(st.session_state["interview_display_index"] + 1) + "/" + str(len(st.session_state["all_interviews"])))
    display_Interview(st.session_state["all_interviews"][st.session_state["interview_display_index"]])


    button_columns=st.columns(5)

    # if button_columns[1].button("Previous"):
    #     if st.session_state["interview_display_index"] == 0:
    #         st.write("No more interviews available.")
    #     else: 
    #         st.session_state["interview_display_index"] -= 1
    #         st.rerun()
    # if button_columns[3].button("Next"):
    #     if st.session_state["interview_display_index"] >= len(st.session_state["all_interviews"])-1:
    #         st.write("No more interviews available")
    #     else: 
    #         st.session_state["interview_display_index"] += 1
    #         st.rerun()

    button_columns[2].button("Back to Login", on_click=set_stage, args=[LOGIN_PAGE])


if st.session_state["stage"] == SETTINGS:    
    st.session_state["interview"] = None
    st.session_state["messages"] = []
    st.session_state["convo_memory"] = []
    st.session_state["convo_summary"]=""
    st.session_state["convo_file"] = None
    st.session_state["convo_prompt"] = ""
    st.session_state["sent"] = False
    st.session_state["start_time"] = date.datetime.now()
    st.session_state["tokens"] = {"convo": {"input": 0, "output": 0},
                                  "class": {"input": 0, "output": 0},
                                  "diag": {"input": 0, "output": 0}}

    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        st.title("Patient Settings")
        st.write("Select your preferred settings for your interview with the virtual patient. Currently we have only made \"John Smith\" available, and voice input is highly encouraged.")
        chat_mode = st.selectbox("Would you like to use text or voice input for the interview?",
                                ["Text", "Voice"],
                                index = None,
                                placeholder = "Select interview mode...")
        if chat_mode == "Text": st.session_state["chat_mode"] = CHAT_INTERFACE_TEXT
        elif chat_mode == "Voice": st.session_state["chat_mode"] = CHAT_INTERFACE_VOICE
        else: st.session_state["chat_mode"] = None

        patient_name = st.selectbox("Which patient would you like to interview?", 
                                    ["John Smith"],
                                    index = None,
                                    placeholder = "Select patient...")
        if patient_name: 
            st.session_state["interview"] = Interview.build(username = st.session_state["username"], 
                                                            patient = Patient.build(patient_name))

        if st.session_state["chat_mode"] and st.session_state["interview"]: 
            st.session_state["sent"]==False
            st.button("Start Interview", on_click=set_stage, args=[CHAT_SETUP])



if st.session_state["stage"] == CHAT_SETUP:
    st.session_state["convo_prompt"] = st.session_state["interview"].patient.convo_prompt
    if(st.session_state["sent"]==False):
        st.session_state["interview"].start_time = str(st.session_state["start_time"])
        st.session_state["session_id"] = hash(st.session_state["interview"].username+st.session_state["interview"].start_time)
        st.session_state["interview"].id=st.session_state["session_id"]
        collection.insert_one(st.session_state["interview"].model_dump())
        st.session_state["sent"]==True

    set_stage(st.session_state["chat_mode"])


if st.session_state["stage"] == CHAT_INTERFACE_TEXT:
    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        st.title("Interview")
        st.write("You may now begin your interview with " + st.session_state["interview"].patient.id + ". Start by introducing yourself.")
        st.write("Click the Restart button to restart the interview. Click the End Interview button to go to the download screen.")

        container = st.container(height=300)

        for message in st.session_state["interview"].messages:
            with container:
                with st.chat_message(message.role):
                    st.markdown(message.content)

        if user_input := st.chat_input("Type here..."):
            with container:
                with st.chat_message("User"):
                    st.markdown(user_input)
            
            response, speech = get_chat_output(user_input)

            with container:
                with st.chat_message("AI"): #TODO Needs avatar eventually
                    st.markdown(response)
                    play_voice(speech)

        columns = st.columns(4)
        columns[1].button("Restart", on_click=set_stage, args=[SETTINGS])
        columns[2].button("End Interview", on_click=set_stage, args=[PHYSICAL_ECG_SCREEN])


if st.session_state["stage"] == CHAT_INTERFACE_VOICE:
    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        st.title("Interview")
        st.write("You may now begin your interview with " + st.session_state["interview"].patient.id + ". Start by introducing yourself.")
        st.write("""Click the Start Recording button to start recording your voice input to the virtual patient.
                The button will then turn into a Stop button, which you can click when you are done talking.
                Click the Restart button to restart the interview, and the End Interview button to go to the download screen.""")

        audio = audiorecorder("Start Recording", "Stop")
        
        container = st.container(height=300)

        for message in st.session_state["interview"].messages:
            with container:
                with st.chat_message(message.role):
                    st.markdown(message.content)

        if len(audio) > 0:
            user_input = transcribe_voice(audio)
            with container:
                with st.chat_message("User"):
                    st.markdown(user_input)
            
            response, speech = get_chat_output(user_input)

            with container:
                with st.chat_message("AI"): # Needs avatar eventually
                    st.markdown(response)
                    play_voice(speech)

        columns = st.columns(4)
        columns[1].button("Restart", on_click=set_stage, args=[SETTINGS])
        columns[2].button("End Interview", on_click=set_stage, args=[PHYSICAL_ECG_SCREEN])


if st.session_state["stage"] == PHYSICAL_ECG_SCREEN:
    st.title("Physical and ECG Information")
    layout1 = st.columns([7, 1])
    layout1[0].write("Now that you've taken a chance to speak with the patient, you can take a chance to take a look at data obtained during a physical upon admittance to the ER. Review the following physical and ECG before proceeding.")
    layout1[1].button("Proceed to Diagnosis", on_click=set_stage, args = [DIAGNOSIS])
    
    # Update the database from before

    collection.replace_one({"username":st.session_state["interview"].username, "start_time":st.session_state["interview"].start_time}, st.session_state["interview"].model_dump())

    layout2 = st.columns([1, 1])
    with layout2[0].container():
        st.header("Physical Examination", divider = "grey")
        physical_exam_doc = Document(st.session_state["interview"].patient.physical)
        for paragraph in physical_exam_doc.paragraphs:
            st.write(paragraph.text)
    with layout2[1].container():
        st.header("ECG", divider = "grey")
        st.image(st.session_state["interview"].patient.ECG)


if st.session_state["stage"] == DIAGNOSIS:
    st.title("Diagnosis")
    st.write("Use the interview transcription and additional patient information (physical examination and ECG) to provide an interpretative summary, a list of potential diagnoses, a rationale reasoning through which diagnoses are more/less likely, and a final diagnosis.")
    st.write("Click the \"Get Feedback\" button once you are all done to automatically receive grading and feedback on your interview and diagnosis.")

    # 2 column full width layout
    layout1 = st.columns([1, 1])

    # User inputs
    summary = layout1[0].text_area(label = "Write an interpretative summary for the patient, recording the key details of the case:", placeholder = "Interpretive summary for patient", height = 200)
    layout11 = layout1[0].columns([1, 1, 1])
    potential1 = layout11[0].text_input(label = "List 3 potential diagnoses:", placeholder = "First condition name")
    potential2 = layout11[1].text_input(label = "None", placeholder = "Second condition name", label_visibility = "hidden")
    potential3 = layout11[2].text_input(label = "None", placeholder = "Third condition name", label_visibility = "hidden")
    rationale = layout1[0].text_area(label = "Write a rationale reasoning through the potential diagnoses you listed in order to determine what is the best diagnosis for the patient:", placeholder = "Rationale for diagnosis")
    final = layout1[0].text_input(label = "Based on your rationale, write your final diagnosis:", placeholder = "Condition name")

    # 3 buttons at bottom
    layout12 = layout1[0].columns([1, 1, 1])
    # New Interview
    # layout12[0].button("New Interview", on_click=set_stage, args=[SETTINGS])
    # Download Interview
    # currentDateAndTime = date.datetime.now()
    # st.session_state["interview"].date_time = str(currentDateAndTime)
    # date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")
    
    bio = io.BytesIO()
    st.session_state["convo_file"] = create_convo_file(st.session_state["interview"].username, 
                                                       st.session_state["interview"].patient.id, 
                                                       [message.model_dump() for message in st.session_state["interview"].messages])
    st.session_state["convo_file"].save(bio)

    date_time = date.datetime.now().strftime("%d-%m-%y__%H-%M")
    layout12[1].download_button("Download interview", 
                                data = bio.getvalue(), 
                                file_name = st.session_state["interview"].username + "_"+date_time + ".docx", 
                                mime = "docx")
    # Get Feedback
    if layout12[2].button("Get Feedback"): 
        st.session_state["interview"].add_diagnosis_inputs(summary, [potential1, potential2, potential3], rationale, final)
        set_stage(FEEDBACK_SETUP)
        st.rerun()
    
    # Interview transcription
    chat_container = layout1[1].container(height=400)
    for message in st.session_state["interview"].messages:
        with chat_container:
            with st.chat_message(message.role):
                st.markdown(message.content)
    # Physical Examination
    with layout1[1].expander("Physical Examination"):
        physical_exam_doc = Document(st.session_state["interview"].patient.physical)
        for paragraph in physical_exam_doc.paragraphs:
            st.write(paragraph.text)
    # ECG
    with layout1[1].expander("ECG"):
        st.image(st.session_state["interview"].patient.ECG)

    
if st.session_state["stage"] == FEEDBACK_SETUP:
    st.title("Processing feedback...")
    st.write("This might take a few minutes.")
    st.session_state["interview"].add_feedback()
    # st.json(st.session_state["interview"].model_dump_json())
    st.session_state["interview_dict"] = st.session_state["interview"].model_dump()
    
    set_stage(FEEDBACK_SCREEN)
    st.rerun()


if st.session_state["stage"] == FEEDBACK_SCREEN:
    st.title("Feedback")
    layout1 = st.columns([7, 1])
    layout1[0].write("This is the WIP (the UI especially) feedback screen. There are 3 tabs of grading/feedback available. \"Data Acquisition\" looks through your interview with the patient and grades you on what information you asked for and what information you obtained from the patient. \"Diagnosis\" is self explanatory, grading you on your summary, diagnoses, and rationale. \"Case Explanation\" provides you with a PDF of a detailed explanation of the John Smith case written by Dr. Corbett.")
    layout1[0].write("Click the \"Go to Survey\" button on the right once you are done looking through the grading/feedback; the next (final) screen will take you through a brief survey on your experience with the application")
    layout1[1].button("Go to Survey", on_click=set_stage, args=[SURVEY])
    
    # Let the display methods cook
    display_Interview(st.session_state["interview"].model_dump())


if st.session_state["stage"] == SURVEY:
    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        st.title("Survey")
        st.write("Please take the time to give us some feedback and constructive criticism so that we can improve this application.")

        #QUESTION 1
        question1="Were there any issues you encountered while interviewing the virtual patient or writing your diagnosis? What could be improved?"
        answer1 = st.text_area(label = question1, height = 100)
        #QUESTION 2
        question2="Were there any issues you encountered in your grading and feedback? What could be improved?"
        answer2 = st.text_area(label = question2, height = 100)
        #QUESTION 3
        question3="Do you have any other suggestions or areas for improvement? This could include ideas for new features, feedback sections, etc."
        answer3 = st.text_area(label = question3, height = 100)
        
        question4="How helpful do you think practicing an interview would be in learning cardiovascular pathophisiology?"
        answer4= st.text_area(label = question4, height = 100)


        #STORING INTERVIEWS IN SEPARATE QUESTIONS DB
        survey["question1"]=[question1,answer1]
        survey["question2"]=[question2,answer2]
        survey["question3"]=[question3,answer3]
        survey["question4"]=[question4,answer4]

        st.session_state["interview"].survey = survey
        columns = st.columns(3)
        columns[1].button("Go to End Screen", on_click=set_stage, args=[FINAL_SCREEN])


if st.session_state["stage"] == FINAL_SCREEN: 
    layout1 = st.columns([1, 3, 1])
    with layout1[1]:
        st.title("Thank you! :heart:")
        st.write("All done! Thank you so much for taking the time to help us test our application. Your interview, diagnosis, and survey has been recorded and sent to us automatically.")
        st.write("Click the download button to download your most recent interview as a word file. Click the New Interview button to go back to the chat interface and keep testing.")
        
        # Record time
        end_time = date.datetime.now()
        time_elapsed = end_time - st.session_state["start_time"]
        st.session_state["interview"].time_elapsed = str(time_elapsed)

        # Record cost
        print(str(st.session_state["tokens"]) + "\n")
        st.session_state["interview"].cost = 0
        st.session_state["interview"].cost += (st.session_state["tokens"]["convo"]["input"] / 1000000) * COSTS[CONVO_MODEL]["input"]
        st.session_state["interview"].cost += (st.session_state["tokens"]["convo"]["output"] / 1000000) * COSTS[CONVO_MODEL]["output"]
        st.session_state["interview"].cost += (st.session_state["tokens"]["class"]["input"] / 1000000) * COSTS[CLASS_MODEL]["input"]
        st.session_state["interview"].cost += (st.session_state["tokens"]["class"]["output"] / 1000000) * COSTS[CLASS_MODEL]["output"]
        st.session_state["interview"].cost += (st.session_state["tokens"]["diag"]["input"] / 1000000) * COSTS[DIAG_MODEL]["input"]
        st.session_state["interview"].cost += (st.session_state["tokens"]["diag"]["output"] / 1000000) * COSTS[DIAG_MODEL]["output"]
        st.write("Estimated cost: $" + str(st.session_state["interview"].cost))


        bio = io.BytesIO()
        st.session_state["convo_file"] = create_convo_file(st.session_state["interview"].username, 
                                                        st.session_state["interview"].patient.id, 
                                                        [message.model_dump() for message in st.session_state["interview"].messages])
        st.session_state["convo_file"].save(bio)
        
        date_time = end_time.strftime("%d-%m-%y__%H-%M")
        button_columns = st.columns(3)
        button_columns[0].download_button("Download interview", 
                        data = bio.getvalue(),
                        file_name = st.session_state["interview"].username + "_"+ date_time + ".docx",
                        mime = "docx")  
        
        # Store interview in database and send email as backup
        if st.session_state["sent"] == False:
            collection.insert_one(st.session_state["interview"].model_dump())
            # send_email(bio, EMAIL_TO_SEND, st.session_state["interview"].username, date_time, None)
            st.session_state["sent"] = True
            
        # st.download_button("Download JSON",
        #             data=st.session_state["interview"].get_json(),
        #             file_name = st.session_state ["interview"].get_username() + "_"+date_time + ".json",
        #             mime="json")

        button_columns[1].button("New Interview", on_click=set_stage, args=[SETTINGS])
        button_columns[2].button("Back to Login", on_click=set_stage, args=[LOGIN_PAGE])