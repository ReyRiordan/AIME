from langchain_community.chat_models import ChatOpenAI
from langchain.chains.conversation.base import ConversationChain
from langchain.memory.buffer import ConversationBufferMemory
from docx import Document
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import (
    Mail, Attachment, FileContent, FileName,
    FileType, Disposition, ContentId)
import os
import datetime as date
import base64
import io
import streamlit as st
from audiorecorder import audiorecorder
from openai import OpenAI
import tempfile
from lookups import *
from website_classes import *
from annotated_text import annotated_text
import json
from LLM_methods import *


def get_webtext(content: str) -> str:
    path = WEBSITE_TEXT[content]
    with open(path, 'r', encoding='utf8') as webtext:
            text = webtext.read()
    return text


def create_convo_file(interview: Interview) -> Document:
    convo = Document()
    heading = convo.add_paragraph("User: " + interview.get_username() + ", ")
    currentDateAndTime = date.datetime.now()
    date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")
    heading.add_run("Date: " + date_time + ", ")
    heading.add_run("Patient: " + interview.get_patient().name)
    for message in interview.get_messages():
        convo.add_paragraph(message.role + ": " + message.content)
    
    return convo


def send_email(bio, EMAIL_TO_SEND, username, date_time, feedback):
    message = Mail(
        from_email = 'rutgers.aime@gmail.com',
        to_emails = EMAIL_TO_SEND,
        subject = "Conversation from " + username + " at time " + date_time,
        html_content = feedback)
    attachment = Attachment()
    encoded = base64.b64encode(bio.getvalue()).decode()
    attachment.file_content=FileContent(encoded)
    attachment.file_type = FileType('docx')
    attachment.file_name = FileName(username + "_" + date_time + ".docx")
    attachment.disposition = Disposition('attachment')
    attachment.content_id = ContentId('docx')
    message.attachment = attachment
    try:
        sg = SendGridAPIClient(os.environ.get('SENDGRID_API_KEY'))
        response = sg.send(message)
        print(response.status_code)
        print(response.body)
        print(response.headers)
    except: 
        print("ERROR ENCOUNTERED SENDING MESSAGE\n")


def display_DataCategory(category: dict[str, str], checklist: dict[str, bool], weights: dict[str, int], score: int, maxscore: int) -> None:
    st.header(f":{category["color"]}[{category["header"]}]: {score}/{maxscore}", divider = category["color"])
    display_labels = [(key, str(weights[key]), "#baffc9" if value else "#ffb3ba") for key, value in checklist.items()]
    annotated_text(display_labels)


# TODO: @Rey since you have an aesthetically pleasing way of doing it please display all of the data from any given Interview instance. This method should take in an Interview and display all relavant data

def display_interview(interview: Interview):
    st.subheader("Username: " + interview.get_username())
    chat_container = st.container(height=300)
    for message in interview.get_messages():
            with chat_container:
                with st.chat_message(message.role):
                    if message.annotation is None:
                        st.markdown(message.content)
                    else:
                        annotated_text((message.content, message.annotation, message.highlight))
    # for category in interview.get_categories():
    #         if category.tab == "data":
    #             display_datagrades(interview.get_datagrades(), category)

def dict_to_patient(json_dict):
    to_return=Patient(json_dict["name"])
    return to_return

def dict_to_interview(json_dict):
    to_return=Interview(json_dict["username"],dict_to_patient(json_dict["patient"]))
    for message_json in json_dict["messages"]:
        to_return.add_message(Message(type=message_json["type"],
                                      role=message_json["role"],
                                      content=message_json["content"]
                                      ))
    return to_return