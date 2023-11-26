from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory

import sys
import os
import time
import textract
import tiktoken
import testing.token_counter as token_counter

# import aspose.words as aw
import datetime as date
from docx import Document
# import export_docx


KEY = ""
MODEL = "gpt-4"
PROMPT = "Prompt/Prompt_10-28.txt"
TOWRITE = "Prompt/output.txt"
QUESTIONS = "Prompt/questions.txt"

llm = ChatOpenAI(openai_api_key=KEY, model_name=MODEL)
conversation = ConversationChain(llm=llm, memory=ConversationBufferMemory())



# A list of all messages sent thusfar
allMessages = []
#Adding the prompt from the .docx file before the conversation begins
#user_input = textract.process("./Prompt/Chat_GPT_prompt_V1023_OBJ1.docx").decode()


# with open(PROMPT, "r", encoding="utf8") as prompt:
#     prompt_input = prompt.read()


if getattr(sys, 'frozen', False):
    with open(os.path.join(sys._MEIPASS, PROMPT), "r", encoding="utf8") as prompt:
        prompt_input=prompt.read()
else:
    with open(PROMPT, "r", encoding="utf8") as prompt:
        prompt_input = prompt.read()


#with open(TOWRITE,"w", encoding = "utf8") as logger: 
    #logger.write(""); 

document = Document()
document.add_heading('Artificial Intelligence for Medical Education')


print("INSTRUCTIONS: Type your answers below. When you are done with the conversation, simply type \"END\" to end the conversation\n\n")
my_name = input("What is your name? ")
document.add_heading(my_name, 2)

output = conversation.predict(input=prompt_input)
# print("GPT: " + output + "\n")
# document.add_paragraph("GPT: "+output)



user_input = input("Begin the conversation: ")
while not "END" in user_input:
    allMessages.append(user_input)
    document.add_paragraph(my_name+": "+user_input)
    output = conversation.predict(input=user_input)
    allMessages.append(output)
    document.add_paragraph("GPT: "+output)
    print("GPT: " + output + "\n")
    user_input = input("User: ")

# Export
# export_docx.exportAsDocx(PROMPT, MODEL, allMessages)

currentDateAndTime = date.datetime.now()
date_time = currentDateAndTime.strftime("%d-%m-%y__%H-%M")
document.save("Medical_Interview_"+date_time+".docx")

#with open(TOWRITE,"a",encoding="utf8") as logger:
    #logger.write(output)
    #logger.write("User: "+user_input+"\n")


